import os
from PyQt6.QtWidgets import QApplication, QMainWindow, QGridLayout, QWidget, QLabel, \
    QPushButton, QStackedWidget, QCalendarWidget, QScrollArea, QVBoxLayout, QTextEdit, QLineEdit, QFileDialog,\
    QPlainTextEdit, QMessageBox, QDialog, QHBoxLayout, QListWidget
import zipfile
from io import BytesIO
import sqlite3
from PyQt6.QtGui import QGuiApplication, QPixmap, QFont, QIcon, QColor
from PyQt6.QtCore import Qt, QDate, QSize
from collections import Counter
from os import remove as remove_file, makedirs, path
from shutil import rmtree
import sys
import requests
from datetime import datetime
from docx import Document

flask_server_url = "http://127.0.0.1:5000/"

authorized = False

font = QFont("Segoe Print", 16)
conn = sqlite3.connect("database.db", check_same_thread=False)
cursor = conn.cursor()

cursor.execute('SELECT * FROM Events;')
events = cursor.fetchall()
events_dates = [QDate.fromString(el[1], 'dd.MM.yyyy') for el in events]
counted_dates = Counter(events_dates)

cursor.execute('SELECT * FROM Organisations;')
orgs = cursor.fetchall()

cursor.execute("SELECT * FROM Projects;")
projects = cursor.fetchall()

cursor.execute("SELECT * FROM Diary;")
notes = sorted(cursor.fetchall(), key=lambda x: datetime.strptime(x[3], '%d.%m.%Y').date())


class SamStack(QStackedWidget):
    def __init__(self, *args):
        super().__init__()
        self.setWindowTitle("Чемоданчик студента")

        for el in args:
            self.addWidget(el)

        self.setMinimumSize(root.width, root.height)
        self.setMaximumSize(int(root.width * 1.1), int(root.height * 1.1))
        self.setWindowIcon(QIcon('suitcase.png'))


class SamRoot(QMainWindow):
    def __init__(self):
        super().__init__()
        self.width = 450
        self.height = 600

        self.setMinimumSize(self.width, self.height)
        self.setMaximumSize(int(self.width * 1.1), int(self.height * 1.1))

        # Создаем QLabel для отображения фонового изображения
        self.background_label = QLabel(self)
        self.background_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.background_label.setScaledContents(True)  # Растягиваем содержимое QLabel

        pixmap = QPixmap("MainBackground.jpg")
        self.background_label.setPixmap(pixmap)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        layout = QGridLayout(self.central_widget)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.background_label, 0, 0, 7, 9)

        desktop_geometry = QGuiApplication.primaryScreen().availableGeometry()
        x = (desktop_geometry.width() - self.width) // 2
        y = (desktop_geometry.height() - self.height) // 2

        self.setGeometry(x, y, self.width, self.height)

        lbl_student = QLabel("Чемоданчик Студента")
        lbl_student.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.calendar = QPushButton("Календарь мероприятий")
        self.world_projects = QPushButton("Мир проектов")
        self.profiles_orgs = QPushButton("Профили организаций")
        self.diary_constructor = QPushButton("Дневник для записей")
        self.options = QPushButton("Настройки")
        self.exit = QPushButton("Выход")

        lbl_student.setFont(QFont("Segoe Print", 18))
        self.calendar.setFont(font)
        self.world_projects.setFont(font)
        self.profiles_orgs.setFont(font)
        self.diary_constructor.setFont(font)
        self.options.setFont(font)
        self.exit.setFont(font)

        lbl_student.setStyleSheet("background-color: rgba(255, 128, 0, 0.5); "
                                  "border-top: 3px solid #ff8000; border-bottom: 3px solid #ff8000;")
        lbl_student.setMaximumHeight(40)

        self.calendar.setStyleSheet("background-color: #bc5600;")
        self.world_projects.setStyleSheet("background-color: #bc5600;")
        self.profiles_orgs.setStyleSheet("background-color: #bc5600;")
        self.diary_constructor.setStyleSheet("background-color: #bc5600;")
        self.options.setStyleSheet("background-color: #bc5600;")
        self.exit.setStyleSheet("background-color: #bc5600;")

        self.calendar.clicked.connect(lambda e: open_another_widget(1,
                                                                 (scalendar.width, scalendar.height),
                                                                 (10 ** 6, 10 ** 6), state="calendar"))
        self.diary_constructor.clicked.connect(lambda e: open_another_widget(2,
                                                                (diary_builder.width, diary_builder.height),
                                                                (10 ** 6, 10 ** 6)))
        self.world_projects.clicked.connect(lambda e: open_another_widget(3,
                                                                (diary_builder.width, diary_builder.height),
                                                                (10 ** 6, 10 ** 6)))
        self.profiles_orgs.clicked.connect(lambda e: open_another_widget(4,
                                                                (diary_builder.width, diary_builder.height),
                                                                (10 ** 6, 10 ** 6)))
        self.options.clicked.connect(lambda e: open_another_widget(5,
                                                                    (settings.width, settings.height),
                                                    (int(settings.width * 1.1), int(settings.height * 1.1))))
        self.exit.clicked.connect(lambda e: app.quit())

        layout.addWidget(lbl_student, 0, 0, 1, 9)
        layout.addWidget(self.calendar, 1, 1, 1, 7)
        layout.addWidget(self.world_projects, 2, 1, 1, 7)
        layout.addWidget(self.diary_constructor, 3, 1, 1, 7)
        layout.addWidget(self.profiles_orgs, 4, 1, 1, 7)
        layout.addWidget(self.options, 5, 1, 1, 7)
        layout.addWidget(self.exit, 6, 1, 1, 7)


class SamLogin(QDialog):
    def __init__(self):
        super().__init__()

        self.setWindowTitle('Авторизация')

        layout = QVBoxLayout(self)

        self.label_username = QLabel('Username:')
        self.edit_username = QLineEdit(self)

        self.label_password = QLabel('Password:')
        self.edit_password = QLineEdit(self)
        self.edit_password.setEchoMode(QLineEdit.EchoMode.Password)  # Скрытие вводимых символов

        self.button_login = QPushButton('Log in', self)
        self.button_login.clicked.connect(self.accept)

        layout.addWidget(self.label_username)
        layout.addWidget(self.edit_username)
        layout.addWidget(self.label_password)
        layout.addWidget(self.edit_password)
        layout.addWidget(self.button_login)


class SamSettings(QWidget):
    def __init__(self):
        super().__init__()
        self.width = 450
        self.height = 350

        self.setMinimumSize(self.width, self.height)
        self.setMaximumSize(int(self.width * 1.1), int(self.height * 1.1))

        # Создаем QLabel для отображения фонового изображения
        self.background_label = QLabel(self)
        self.background_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.background_label.setScaledContents(True)  # Растягиваем содержимое QLabel

        pixmap = QPixmap("MainBackground.jpg")
        self.background_label.setPixmap(pixmap)

        layout = QGridLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        desktop_geometry = QGuiApplication.primaryScreen().availableGeometry()
        x = (desktop_geometry.width() - self.width) // 2
        y = (desktop_geometry.height() - self.height) // 2

        self.setGeometry(x, y, self.width, self.height)

        lbl_student = QLabel("Настройки")
        lbl_student.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.auth = QPushButton("Авторизироваться")
        self.reload = QPushButton("Обновить данные")
        self.btn_clear_all = QPushButton("Очистить хранилище")
        self.back_btn = QPushButton("Назад")

        lbl_student.setFont(QFont("Segoe Print", 20))
        self.auth.setFont(font)
        self.reload.setFont(font)
        self.btn_clear_all.setFont(font)
        self.back_btn.setFont(font)

        lbl_student.setStyleSheet("color: #e00000;")
        lbl_student.setMaximumHeight(40)

        self.auth.setStyleSheet("background-color: #bc5600;")
        self.reload.setStyleSheet("background-color: #bc5600;")
        self.btn_clear_all.setStyleSheet("background-color: #bc5600;")
        self.back_btn.setStyleSheet("background-color: #bc5600;")

        self.auth.clicked.connect(self.authorize)
        self.reload.clicked.connect(self.load_data)
        self.back_btn.clicked.connect(lambda e: open_another_widget(0, (root.width, root.height),
                                                                    (int(root.width * 1.1), int(root.height * 1.1))))
        self.btn_clear_all.clicked.connect(self.clear_all)

        layout.addWidget(self.background_label, 0, 0, 5, 9)
        layout.addWidget(lbl_student, 0, 0, 1, 9)
        layout.addWidget(self.auth, 1, 1, 1, 7)
        layout.addWidget(self.reload, 2, 1, 1, 7)
        layout.addWidget(self.btn_clear_all, 3, 1, 1, 7)
        layout.addWidget(self.back_btn, 4, 1, 1, 7)

        self.login = SamLogin()

    def clear_all(self):
        global events_dates, events, scalendar, counted_dates, profiles_orgs, orgs, stack_root, projects, world_projects
        try:
            rmtree("Prof_pictures")
            rmtree("Orgs_pictures")
            makedirs("Orgs_pictures")
            makedirs("Prof_pictures")
            cursor.execute("DELETE FROM Events;")
            cursor.execute("DELETE FROM Organisations;")
            cursor.execute("DELETE FROM Projects;")
            conn.commit()

            events_dates = []
            events = []
            counted_dates = Counter()
            stack_root.removeWidget(scalendar)
            del scalendar
            scalendar = SamCalendar()
            stack_root.insertWidget(1, scalendar)

            orgs = []
            stack_root.removeWidget(profiles_orgs)
            del profiles_orgs
            profiles_orgs = SamProfilesOrgs()
            stack_root.insertWidget(4, profiles_orgs)

            projects = []
            stack_root.removeWidget(world_projects)
            del world_projects
            world_projects = SamWorldProjects()
            stack_root.insertWidget(3, world_projects)

            QMessageBox.information(self, "Очистка", "Очистка прошла успешно.")
        except Exception as e:
            QMessageBox.critical(self, "Очистка", "Ошибка при очистке данных...")

    def load_data(self):
        response = None
        try:
            response = requests.get(flask_server_url + "download_all_files")
            response.raise_for_status()  # Проверка успешности запроса
        except requests.exceptions.RequestException as e:
            QMessageBox.critical(self, "Обновление базы данных", "Ошибка при подключении к серверу.")
            return

        if response.status_code == 200:
            global conn, cursor, events, events_dates, counted_dates, scalendar, profiles_orgs, orgs, stack_root, \
                    world_projects, projects
            # Создаем временный буфер для хранения данных архива
            zip_buffer = BytesIO(response.content)

            # Распаковываем архив в указанную директорию
            with zipfile.ZipFile(zip_buffer, 'r') as zip_ref:
                zip_ref.extractall(".")

            conn2 = sqlite3.connect("database_server.db", check_same_thread=False)
            cursor2 = conn2.cursor()

            copy_unique_records(cursor2, cursor, "Events", ["Date", "Name", "Description", "Pictures"])
            copy_unique_records(cursor2, cursor, "Organisations", ["Name", "Mission", "Contact",
                                                                   "Category", "Projects", "Photo"])
            copy_unique_records(cursor2, cursor, "Projects", ["Name", "Description",
                                                              "Photo", "Category", "Owner"])

            cursor2.close()
            conn2.close()
            remove_file("database_server.db")

            cursor.execute('SELECT * FROM Events;')
            events = cursor.fetchall()
            events_dates = [QDate.fromString(el[1], 'dd.MM.yyyy') for el in events]
            counted_dates = Counter(events_dates)

            stack_root.removeWidget(scalendar)
            del scalendar
            scalendar = SamCalendar()
            stack_root.insertWidget(1, scalendar)

            cursor.execute('SELECT * FROM Organisations;')
            orgs = cursor.fetchall()

            stack_root.removeWidget(profiles_orgs)
            del profiles_orgs
            profiles_orgs = SamProfilesOrgs()
            stack_root.insertWidget(4, profiles_orgs)

            cursor.execute("SELECT * FROM Projects;")
            projects = cursor.fetchall()
            stack_root.removeWidget(world_projects)
            del world_projects
            world_projects = SamWorldProjects()
            stack_root.insertWidget(3, world_projects)

            QMessageBox.information(self, "Обновление данных", "Архив успешно получен и распакован.")
        else:
            QMessageBox.critical(self, "Обновление базы данных",
                                       f"Ошибка при получении архива. Код состояния: {response.status_code}")

    def authorize(self):
        global authorized
        result = self.login.exec()
        username = self.login.edit_username.text()
        password = self.login.edit_password.text()
        self.login.edit_password.setText("")
        self.login.edit_username.setText("")
        if result == QDialog.DialogCode.Accepted:
            try:
                response = requests.get(flask_server_url + "auth", params={"username": username, "password": password})
                response.raise_for_status()  # Проверка успешности запроса
            except requests.exceptions.RequestException as e:
                QMessageBox.critical(self, "Авторизация", "Ошибка при подключении к серверу.")
                return

            if response.status_code == 200:
                json_data = response.json()
                if json_data:
                    if json_data["status"] == "OK":
                        authorized = (username, password)
                        QMessageBox.information(self, "Авторизация", "Авторизация прошла успешно.")
                    elif json_data["status"] == "Error":
                        QMessageBox.warning(self, "Авторизация", "Неверное имя пользователя или пароль.")
                    else:
                        QMessageBox.critical(self, "Авторизация", "Неизвестная ошибка.")
            else:
                QMessageBox.critical(self, "Авторизация",
                                     f"Ошибка при попытке авторизации. Код состояния: {response.status_code}")


def copy_unique_records(cursor_source, cursor_destination, table_name, cols):
    cursor_source.execute(f"SELECT * FROM {table_name};")
    data_to_merge1 = [tuple(data[1:]) for data in cursor_source.fetchall()]

    cursor_destination.execute(f"SELECT * FROM {table_name};")
    data_to_merge2 = [tuple(data[1:]) for data in cursor_destination.fetchall()]

    data = set(data_to_merge1 + data_to_merge2)

    cursor_destination.execute(f"DELETE FROM {table_name};")

    i = 1
    for el in data:
        cursor_destination.execute(f'INSERT INTO {table_name} ("Index", {", ".join(cols)}) VALUES (?, '
                                   f'{", ".join(["?" for _ in range(len(cols))])});',
                                   (tuple([i] + [el[i] for i in range(len(cols))])))
        i += 1

    # Фиксация изменений
    cursor_destination.connection.commit()


class SamCalendar(QWidget):
    def __init__(self):
        super().__init__()
        self.width = 1200
        self.height = 700

        self.setMinimumSize(self.width, self.height)

        lbl = QLabel(self)
        pixmap = QPixmap("Wood.jpg")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setScaledContents(True)  # Растягиваем содержимое QLabel

        lbl.setPixmap(pixmap)

        self.back_btn = QPushButton("Назад", self)
        self.back_btn.setFont(font)
        self.back_btn.clicked.connect(lambda e: open_another_widget(0, (root.width, root.height),
                                                            (int(root.width * 1.1), int(root.height * 1.1))))
        self.back_btn.setMinimumHeight(35)
        self.back_btn.setMaximumHeight(45)
        self.back_btn.setStyleSheet("background-color: orange;")

        self.main_calendar = QCalendarWidget()
        self.main_calendar.showToday()
        self.main_calendar.setGridVisible(True)
        self.reload_cdr()

        self.main_calendar.selectionChanged.connect(self.check_date)

        # Создаем QScrollArea
        scroll_area = QScrollArea(self)
        scroll_area.setWidgetResizable(True)  # Разрешаем изменение размера виджета внутри QScrollArea

        # Создаем виджет, который будет помещен в QScrollArea
        scroll_content = QWidget(scroll_area)
        self.layout2 = QVBoxLayout(scroll_content)

        lbl2 = QLabel("Выберите\nДату", self)
        lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl2.setMinimumHeight(180)
        lbl2.setFont(QFont("Segoe Print", 40))
        lbl2.setStyleSheet("border: none; color: blue;")
        self.layout2.addWidget(lbl2)

        # Устанавливаем виджет в QScrollArea
        scroll_area.setWidget(scroll_content)

        layout = QGridLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        layout.addWidget(lbl, 0, 0, 25, 25)
        layout.addWidget(self.main_calendar, 1, 1, 12, 10)
        layout.addWidget(self.back_btn, 23, 0, 2, 5)
        layout.addWidget(scroll_area, 1, 11, 23, 13)
        self.setLayout(layout)

    def reload_cdr(self):
        if events_dates:
            for sdate in events_dates:
                if counted_dates[sdate] == 1:
                    self.main_calendar.setDateTextFormat(sdate, self.get_marked_format((40, 250, 40)))
                elif counted_dates[sdate] == 2:
                    self.main_calendar.setDateTextFormat(sdate, self.get_marked_format((250, 250, 10)))
                else:
                    self.main_calendar.setDateTextFormat(sdate, self.get_marked_format((250, 10, 10)))

    def check_date(self):
        selected_date = self.main_calendar.selectedDate()
        indexes = []

        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        for i in range(len(events_dates)):
            if events_dates[i].toPyDate() == selected_date.toPyDate():
                indexes.append(i)

        lbls = []
        if indexes:
            for i in indexes:
                name = QTextEdit(self)
                name.setText(events[i][2])
                name.setFont(font)
                name.setMinimumHeight(85)
                name.setReadOnly(True)
                name.setStyleSheet("border: none; background-color: #e080e0;")
                lbls.append(name)

                if events[i][4]:
                    lbl = QLabel(self)
                    pixmap = QPixmap("Orgs_pictures/" + events[i][4])

                    lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    lbl.setScaledContents(True)  # Растягиваем содержимое QLabel
                    lbl.setPixmap(pixmap)
                    lbl.setMaximumSize(574, int(pixmap.height() * 574 / pixmap.width()))
                    lbls.append(lbl)

                text = QPlainTextEdit(self)
                text.setPlainText(events[i][3])
                text.setFont(QFont("Segoe Print", 12))
                text.setReadOnly(True)
                text.setStyleSheet("border: none;")
                text.setMinimumHeight(len(text.toPlainText()) + 75)
                lbls.append(text)

            for el in range(len(lbls)):
                self.layout2.addWidget(lbls[el])
        else:
            lbl2 = QLabel("Пусто", self)
            lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
            lbl2.setMinimumHeight(180)
            lbl2.setFont(QFont("Segoe Print", 40))
            lbl2.setStyleSheet("border: none; color: blue;")
            self.layout2.addWidget(lbl2)

    def get_marked_format(self, color):
        # Установка стиля для помеченных дат (синий цвет)
        marked_format = self.main_calendar.dateTextFormat(QDate.currentDate())
        marked_format.setBackground(QColor(color[0], color[1], color[2]))
        marked_format.setForeground(QColor(0, 0, 0))

        return marked_format


class SamDiaryBuilder(QWidget):
    def __init__(self):
        super().__init__()
        self.width = 1200
        self.height = 700

        self.setMinimumSize(self.width, self.height)

        lbl = QLabel(self)
        pixmap = QPixmap("Paper.jpg")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setScaledContents(True)  # Растягиваем содержимое QLabel

        lbl.setPixmap(pixmap)

        self.back_btn = QPushButton("Назад", self)
        self.back_btn.setFont(font)
        self.back_btn.clicked.connect(lambda e: open_another_widget(0, (root.width, root.height),
                                                                    (int(root.width * 1.1), int(root.height * 1.1))))
        self.back_btn.setMinimumHeight(35)
        self.back_btn.setMaximumHeight(45)
        self.back_btn.setStyleSheet("background-color: orange;")

        self.layout = QGridLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)

        self.for_text = QTextEdit(self)
        self.for_text.setFont(QFont("Segoe Print", 13))
        self.for_text.setPlaceholderText("Пишите здесь...")
        self.for_text.setStyleSheet("background-color: #d0d0d0")

        self.btn_save_it = QPushButton("Сохранить", self)
        self.btn_save_it.setMinimumHeight(35)
        self.btn_save_it.setMaximumHeight(45)
        self.btn_save_it.setStyleSheet("background-color: orange;")
        self.btn_save_it.setFont(font)
        self.btn_save_it.clicked.connect(self.save_note)

        self.btn_save_file = QPushButton("В файл", self)
        self.btn_save_file.setMinimumHeight(35)
        self.btn_save_file.setMaximumHeight(45)
        self.btn_save_file.setStyleSheet("background-color: orange;")
        self.btn_save_file.setFont(font)
        self.btn_save_file.clicked.connect(self.save_note_file)

        self.name = QLineEdit(self)
        self.name.setMinimumHeight(35)
        self.name.setMaximumHeight(45)
        self.name.setStyleSheet("background-color: orange;")
        self.name.setFont(font)
        self.name.setPlaceholderText("Название...")
        self.name.setMaxLength(25)

        scroll_area = QScrollArea(self)
        scroll_area.setStyleSheet("background-color: #d0d0d0")
        scroll_area.setWidgetResizable(True)  # Разрешаем изменение размера виджета внутри QScrollArea

        # Создаем виджет, который будет помещен в QScrollArea
        scroll_content = QWidget(scroll_area)
        self.layout2 = QVBoxLayout(scroll_content)

        self.lst_wdgts = dict()
        for i in range(len(notes)):
            name = QPushButton(self)
            name.setText(notes[i][1] + "\n" + notes[i][3])
            name.setFont(QFont("Segoe Print", 10))
            name.setMaximumSize(183, 40)
            name.setStyleSheet("background-color: #a0a0a0")
            self.lst_wdgts[name] = notes[i][0]
            name.clicked.connect(self.show_history)

        for el in self.lst_wdgts.keys():
            self.layout2.addWidget(el)

        # Устанавливаем виджет в QScrollArea
        scroll_area.setWidget(scroll_content)

        lbl2 = QLabel("Записи")
        lbl2.setFont(QFont("Segoe Print", 18))
        lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl2.setStyleSheet("background-color: rgba(255, 128, 0, 0.7);")

        self.cancel = None
        self.delete_btn = None

        self.layout.addWidget(lbl, 0, 0, 16, 16)
        self.layout.addWidget(lbl2, 0, 0, 1, 3)
        self.layout.addWidget(scroll_area, 1, 0, 12, 3)
        self.layout.addWidget(self.for_text, 1, 4, 12, 11)
        self.layout.addWidget(self.name, 0, 4, 1, 5)
        self.layout.addWidget(self.btn_save_it, 13, 13, 1, 2)
        self.layout.addWidget(self.btn_save_file, 13, 11, 1, 2)
        self.layout.addWidget(self.back_btn, 15, 0, 1, 3)
        self.setLayout(self.layout)

        self.changed_index = None
        self.counter = 0

    def reload_notes(self):
        global notes
        cursor.execute("SELECT * FROM Diary;")
        notes = sorted(cursor.fetchall(), key=lambda x: datetime.strptime(x[3], '%d.%m.%Y').date())
        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        self.lst_wdgts = dict()
        for i in range(len(notes)):
            name = QPushButton(self)
            name.setText(notes[i][1] + "\n" + notes[i][3])
            name.setFont(QFont("Segoe Print", 10))
            name.setMaximumSize(183, 40)
            name.setStyleSheet("background-color: #a0a0a0")
            self.lst_wdgts[name] = notes[i][0]
            name.clicked.connect(self.show_history)

        for el in self.lst_wdgts.keys():
            self.layout2.addWidget(el)

    def save_note_file(self):
        if self.name.text() == "" and self.for_text.toPlainText() == "":
            QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле названия и текста.")
        elif self.name.text() == "":
            QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле названия.")
        elif self.for_text.toPlainText() == "":
            QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле текста.")
        else:
            path, b_p = QFileDialog.getSaveFileName(self, "Сохранить файл", "C:\\", "Word Documents (*.docx);")
            if b_p:
                doc = Document()
                # Добавляем текст в документ
                doc.add_paragraph(self.name.text())
                doc.add_paragraph(self.for_text.toPlainText())
                # Сохраняем документ в указанный файл
                doc.save(path)
                QMessageBox.information(self, "Операция завершена", "Запись была сохранена.")
                self.cancel_it()

    def delete_note(self):
        # Выполнение запроса с передачей индекса записи для удаления
        cursor.execute('DELETE FROM Diary WHERE "Index" = ?;', (self.changed_index,))

        # Фиксация изменений в базе данных
        conn.commit()
        self.cancel_it()
        self.reload_notes()
        QMessageBox.information(self, "Удаление", "Операция завершена успешно..")

    def save_note(self):
        if self.changed_index:
            cursor.execute('UPDATE Diary SET Name = ?, Description = ? WHERE "Index" = ?;',
                           (self.name.text(), self.for_text.toPlainText(), self.changed_index))
            # фиксация изменений в базе данных
            conn.commit()
            self.changed_index = None
            self.cancel_it()
        else:
            cursor.execute('SELECT MAX("Index") FROM Diary;')
            max_index = cursor.fetchone()[0]
            if self.name.text() == "" and self.for_text.toPlainText() == "":
                QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле названия и текста.")
            elif self.name.text() == "":
                QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле названия.")
            elif self.for_text.toPlainText() == "":
                QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле текста.")
            else:
                cursor.execute('INSERT INTO Diary ("Index", Name, Description, Date) VALUES (?, ?, ?, ?);',
                               (max_index + 1, self.name.text(),
                                self.for_text.toPlainText(), QDate.currentDate().toString("dd.MM.yyyy")))
                # фиксация изменений в базе данных
                conn.commit()
                QMessageBox.information(self, "Операция завершена", "Запись была сохранена.")
                self.cancel_it()
                self.reload_notes()

    def show_history(self):
        for el in self.lst_wdgts:
            el.setStyleSheet("background-color: #a0a0a0;")
        sup_el = self.sender()
        sup_el.setStyleSheet("background-color: orange;")
        self.changed_index = self.lst_wdgts[sup_el]
        cursor.execute('SELECT * FROM Diary WHERE "Index" = ?;', (int(self.changed_index),))
        data = cursor.fetchone()
        self.name.setText(data[1])
        self.for_text.setText(data[2])

        if not self.counter:
            self.btn_save_it.setText("Изменить")
            self.cancel = QPushButton("Отменить", self)
            self.cancel.setMinimumHeight(35)
            self.cancel.setMaximumHeight(45)
            self.cancel.setStyleSheet("background-color: orange;")
            self.cancel.setFont(font)
            self.cancel.clicked.connect(self.cancel_it)

            self.delete_btn = QPushButton("Удалить", self)
            self.delete_btn.setMinimumHeight(35)
            self.delete_btn.setMaximumHeight(45)
            self.delete_btn.setStyleSheet("background-color: orange;")
            self.delete_btn.setFont(font)
            self.delete_btn.clicked.connect(self.delete_note)

            self.layout.addWidget(self.cancel, 13, 4, 1, 2)
            self.layout.addWidget(self.delete_btn, 13, 7, 1, 2)
            self.counter += 2

    def cancel_it(self):
        self.changed_index = None
        while self.counter:
            item = self.layout.takeAt(self.layout.count() - 1)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            self.counter -= 1
        self.name.setText("")
        self.for_text.setText("")
        self.btn_save_it.setText("Сохранить")
        for el in self.lst_wdgts:
            el.setStyleSheet("background-color: #a0a0a0;")


class SamNewProj(QWidget):
    def __init__(self):
        super().__init__()
        self.width = 1200
        self.height = 700

        self.setMinimumSize(self.width, self.height)

        lbl = QLabel(self)
        pixmap = QPixmap("Projects.jpg")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setScaledContents(True)  # Растягиваем содержимое QLabel

        lbl.setPixmap(pixmap)

        self.back_btn = QPushButton("Назад", self)
        self.back_btn.setFont(font)
        self.back_btn.clicked.connect(lambda e: open_another_widget(3,
                                               (diary_builder.width, diary_builder.height),
                                               (10 ** 6, 10 ** 6)))
        self.back_btn.setMinimumHeight(35)
        self.back_btn.setMaximumHeight(45)
        self.back_btn.setStyleSheet("background-color: #d77d31;")

        self.layout = QGridLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)

        self.for_text = QTextEdit(self)
        self.for_text.setFont(QFont("Segoe Print", 13))
        self.for_text.setPlaceholderText("Пишите здесь...")
        self.for_text.setStyleSheet("background-color: #d0d0d0")

        self.btn_save_it = QPushButton("Сохранить", self)
        self.btn_save_it.setMinimumHeight(35)
        self.btn_save_it.setMaximumHeight(45)
        self.btn_save_it.setStyleSheet("background-color: orange;")
        self.btn_save_it.setFont(font)
        self.btn_save_it.clicked.connect(self.create_it)

        self.name = QLineEdit(self)
        self.name.setMinimumHeight(35)
        self.name.setMaximumHeight(45)
        self.name.setStyleSheet("background-color: orange;")
        self.name.setFont(font)
        self.name.setPlaceholderText("Название...")

        self.pict = QPushButton("Добавить картинку", self)
        self.pict.setMinimumHeight(35)
        self.pict.setMaximumHeight(45)
        self.pict.setStyleSheet("background-color: orange;")
        self.pict.setFont(font)
        self.pict.clicked.connect(self.choose)

        self.lst = QListWidget(self)
        self.lst.setMaximumSize(300, 250)
        self.lst.setFont(QFont("Segoe Print", 13))
        self.lst.setStyleSheet("background-color: #d0d0d0")

        self.lst.addItem("IT")
        self.lst.addItem("Engineering")
        self.lst.addItem("Social")
        self.lst.addItem("Other")

        lbl2 = QLabel("Новый проект")
        lbl2.setFont(QFont("Segoe Print", 18))
        lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl2.setStyleSheet("background-color: rgba(255, 128, 0, 0.7);")

        lbl3 = QLabel("Категории")
        lbl3.setFont(QFont("Segoe Print", 18))
        lbl3.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl3.setStyleSheet("background-color: rgba(255, 168, 100, 0.8);")

        self.cancel = None
        self.delete_btn = None

        self.layout.addWidget(lbl, 0, 0, 16, 16)
        self.layout.addWidget(lbl2, 0, 7, 1, 5)
        self.layout.addWidget(self.for_text, 1, 1, 12, 11)
        self.layout.addWidget(self.name, 0, 1, 1, 5)
        self.layout.addWidget(self.btn_save_it, 13, 10, 1, 2)
        self.layout.addWidget(self.pict, 13, 1, 1, 4)
        self.layout.addWidget(lbl3, 1, 12, 1, 4)
        self.layout.addWidget(self.lst, 2, 12, 6, 4)
        self.layout.addWidget(self.back_btn, 15, 0, 1, 3)
        self.setLayout(self.layout)
        self.chosen_file = False

    def choose(self):
        path, b_p = QFileDialog.getOpenFileName(self, "Выбрать картинку", "C:\\", "Images (*.png *.xpm *.jpg *.bmp);")
        if b_p:
            self.chosen_file = path

    def create_it(self):
        global authorized
        if self.name.text() == "" and self.for_text.toPlainText() == "":
            QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле названия и текста.")
        elif self.name.text() == "":
            QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле названия.")
        elif self.for_text.toPlainText() == "":
            QMessageBox.critical(self, "Неполные данные", "Вы не заполнили поле текста.")
        elif self.lst.currentItem() is None:
            QMessageBox.critical(self, "Неполные данные", "Вы не указали категорию проекта.")
        else:
            if self.chosen_file:
                try:
                    response = requests.post(flask_server_url + "add_proj", data={"name": self.name.text(),
                                                                     "text": self.for_text.toPlainText(),
                                                                     "category": self.lst.currentItem().text(),
                                                                     "username": authorized[0]},
                                                                     files={'file': (path.basename(self.chosen_file),
                                                                                     open(self.chosen_file, 'rb'))})
                    response.raise_for_status()  # Проверка успешности запроса
                except requests.exceptions.RequestException as e:
                    QMessageBox.critical(self, "Запрос", "Ошибка при подключении к серверу.")
                    return
            else:
                try:
                    response = requests.post(flask_server_url + "add_proj", data={"name": self.name.text(),
                                                                     "text": self.for_text.toPlainText(),
                                                                     "category": self.lst.currentItem().text(),
                                                                     "username": authorized[0]})
                    response.raise_for_status()  # Проверка успешности запроса
                except requests.exceptions.RequestException as e:
                    QMessageBox.critical(self, "Запрос", "Ошибка при подключении к серверу.")
                    return

            if response.status_code == 200:
                json_data = response.json()
                if json_data:
                    if json_data["status"] == "OK":
                        QMessageBox.information(self, "Запрос", "Запрос на создание проекта был отправлен. "
                                                                "Обновите базу данных.")
                    elif json_data["status"] == "error":
                        QMessageBox.warning(self, "Запрос", "Произошла ошибка.")
                    else:
                        QMessageBox.critical(self, "Запрос", "Неизвестная ошибка.")
            else:
                QMessageBox.critical(self, "Запрос", f"Ошибка. Код состояния: {response.status_code}")
            self.chosen_file = False
            self.name.setText("")
            self.for_text.setText("")
            self.lst.clearSelection()


class SamDelProj(QWidget):
    def __init__(self, elst):
        super().__init__()
        self.width = 450
        self.height = 600

        self.setMinimumSize(self.width, self.height)

        lbl = QLabel(self)
        pixmap = QPixmap("Gradient.jpg")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setScaledContents(True)  # Растягиваем содержимое QLabel

        lbl.setPixmap(pixmap)

        self.back_btn = QPushButton("Назад", self)
        self.back_btn.setFont(font)
        self.back_btn.clicked.connect(lambda e: open_another_widget(3,
                                               (diary_builder.width, diary_builder.height),
                                               (10 ** 6, 10 ** 6)))
        self.back_btn.setMinimumHeight(35)
        self.back_btn.setMaximumHeight(45)
        self.back_btn.setStyleSheet("background-color: orange;")

        self.layout = QGridLayout(self)
        self.layout.setContentsMargins(0, 0, 0, 0)

        self.btn_delete_it = QPushButton("Удалить", self)
        self.btn_delete_it.setMinimumHeight(35)
        self.btn_delete_it.setMaximumHeight(45)
        self.btn_delete_it.setStyleSheet("background-color: orange;")
        self.btn_delete_it.setFont(font)
        self.btn_delete_it.clicked.connect(self.erased_one)

        self.lst = QListWidget(self)
        self.lst.setMaximumHeight(500)
        self.lst.setFont(QFont("Segoe Print", 16))
        self.lst.setStyleSheet("background-color: #d0d0d0")

        for i in range(len(elst)):
            self.lst.addItem(elst[i][0])

        lbl2 = QLabel("Удаление проектов", self)
        lbl2.setMinimumHeight(35)
        lbl2.setMaximumHeight(45)
        lbl2.setStyleSheet("color: red;")
        lbl2.setFont(QFont("Segoe Print", 25))
        lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.cancel = None
        self.delete_btn = None

        self.layout.addWidget(lbl, 0, 0, 7, 9)
        self.layout.addWidget(lbl2, 0, 0, 1, 9)
        self.layout.addWidget(self.lst, 1, 0, 5, 9)
        self.layout.addWidget(self.back_btn, 6, 0, 1, 4)
        self.layout.addWidget(self.btn_delete_it, 6, 5, 1, 4)
        self.setLayout(self.layout)

    def erased_one(self):
        if self.lst.currentItem():
            global authorized, del_proj_w, stack_root
            name = self.lst.currentItem().text()

            try:
                response = requests.post(flask_server_url + "delete_proj", data={"username": authorized[0],
                                                                                 "name": name})
                response.raise_for_status()  # Проверка успешности запроса
            except requests.exceptions.RequestException as e:
                QMessageBox.critical(self, "Запрос", "Ошибка при подключении к серверу.")
                return

            if response.status_code == 200:
                json_data = response.json()
                if json_data:
                    if json_data["status"] == "OK":
                        QMessageBox.information(self, "Запрос", "Запрос на удаление был отправлен. "
                                                                "Обновите базу данных в настройках.")
                        cursor.execute("SELECT Photo FROM Projects WHERE Owner = ? AND Name = ?;",
                                       (authorized[0], name))
                        file = cursor.fetchone()
                        if file[0]:
                            os.remove("Proj_pictures/" + file)

                        cursor.execute('DELETE FROM Projects WHERE Owner = ? AND Name = ?;',
                                       (authorized[0], name))
                        cursor.connection.commit()

                        try:
                            response = requests.get(flask_server_url + "get_user_projs",
                                                    data={"username": authorized[0]})
                            response.raise_for_status()  # Проверка успешности запроса
                        except requests.exceptions.RequestException as e:
                            QMessageBox.critical(self, "Запрос", "Ошибка при подключении к серверу.")
                            return

                        if response.status_code == 200:
                            json_data = response.json()
                            if json_data:
                                if json_data["status"] != "OK":
                                    QMessageBox.critical(self, "Запрос", "Неизвестная ошибка.")
                                    return
                        else:
                            QMessageBox.critical(self, "Запрос", f"Ошибка. Код состояния: {response.status_code}")
                            return
                        result = json_data["result"]
                        stack_root.removeWidget(del_proj_w)
                        del del_proj_w
                        del_proj_w = SamDelProj(result)
                        stack_root.insertWidget(7, del_proj_w)
                        stack_root.setCurrentIndex(7)
                    else:
                        QMessageBox.critical(self, "Запрос", "Неизвестная ошибка.")
            else:
                QMessageBox.critical(self, "Запрос", f"Ошибка. Код состояния: {response.status_code}")

        else:
            QMessageBox.critical(self, "Удаление проекта", "Вы не указали проект для удаления.")


class SamWorldProjects(QWidget):
    def __init__(self):
        super().__init__()
        self.width = 1200
        self.height = 700

        self.setMinimumSize(self.width, self.height)

        lbl = QLabel(self)
        pixmap = QPixmap("Projects.jpg")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setScaledContents(True)  # Растягиваем содержимое QLabel

        lbl.setPixmap(pixmap)

        self.back_btn = QPushButton("Назад", self)
        self.back_btn.setMinimumHeight(35)
        self.back_btn.setMaximumHeight(45)
        self.back_btn.setStyleSheet("background-color: orange;")
        self.back_btn.setFont(font)
        self.back_btn.clicked.connect(lambda e: open_another_widget(0, (root.width, root.height),
                                                                    (int(root.width * 1.1), int(root.height * 1.1))))

        layout = QGridLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # Создаем QScrollArea
        scroll_area = QScrollArea(self)
        scroll_area.setStyleSheet("background-color: #d0d0d0")
        scroll_area.setWidgetResizable(True)  # Разрешаем изменение размера виджета внутри QScrollArea

        # Создаем виджет, который будет помещен в QScrollArea
        scroll_content = QWidget(scroll_area)
        self.layout2 = QVBoxLayout(scroll_content)

        self.btn_names = dict()

        bird_pict = QIcon(QPixmap("Bird.png"))
        lst_wdgts = []
        for i in range(len(projects)):
            name = QLineEdit(self)
            name.setText(projects[i][1])
            name.setFont(font)
            name.setMinimumHeight(35)
            name.setMaximumSize(800, 45)
            name.setReadOnly(True)
            name.setStyleSheet("background-color: #ffd550")

            bird = QPushButton(self)
            bird.setIcon(bird_pict)
            bird.setFixedSize(35, 35)
            bird.setIconSize(QSize(30, 30))
            bird.clicked.connect(self.fly_bird)
            self.btn_names[bird] = name.text()

            wdgt = QWidget(self)
            wdgt.setMaximumHeight(60)
            wdgt.setMinimumHeight(45)
            lyt = QHBoxLayout(wdgt)
            lyt.addWidget(name)
            lyt.addWidget(bird)
            lst_wdgts.append(wdgt)

            if projects[i][3]:
                lbl2 = QLabel(self)
                pixmap = QPixmap("Proj_pictures/" + projects[i][3])

                lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                lbl2.setPixmap(pixmap)
                lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                lst_wdgts.append(lbl2)

            text = QTextEdit(self)
            text.setPlainText(projects[i][2])
            text.setFont(font)
            text.setMinimumHeight(int(len(projects[i][2]) * 1.1))
            text.setReadOnly(True)
            lst_wdgts.append(text)

        for i in range(len(lst_wdgts)):
            self.layout2.addWidget(lst_wdgts[i])

        # Устанавливаем виджет в QScrollArea
        scroll_area.setWidget(scroll_content)

        lbl3 = QLabel("Категории", self)
        lbl3.setMinimumHeight(35)
        lbl3.setMaximumHeight(45)
        lbl3.setStyleSheet("color: red;")
        lbl3.setFont(QFont("Segoe Print", 25))
        lbl3.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.btn_all = QPushButton("Все", self)
        self.btn_all.setStyleSheet("background-color: orange;")
        self.btn_all.setMinimumHeight(40)
        self.btn_all.setFont(font)
        self.btn_all.clicked.connect(self.activate_all)

        self.btn_it = QPushButton("IT", self)
        self.btn_it.setStyleSheet("background-color: #d0d0d0;")
        self.btn_it.setMinimumHeight(40)
        self.btn_it.setFont(font)
        self.btn_it.clicked.connect(lambda e: self.activate_something(self.btn_it, "IT"))

        self.btn_eng = QPushButton("Инженерия", self)
        self.btn_eng.setStyleSheet("background-color: #d0d0d0;")
        self.btn_eng.setMinimumHeight(40)
        self.btn_eng.setFont(font)
        self.btn_eng.clicked.connect(lambda e: self.activate_something(self.btn_eng, "Engineering"))

        self.btn_social = QPushButton("Социум", self)
        self.btn_social.setStyleSheet("background-color: #d0d0d0;")
        self.btn_social.setMinimumHeight(40)
        self.btn_social.setFont(font)
        self.btn_social.clicked.connect(lambda e: self.activate_something(self.btn_social, "Social"))

        self.btn_other = QPushButton("Другое", self)
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setMinimumHeight(40)
        self.btn_other.setFont(font)
        self.btn_other.clicked.connect(lambda e: self.activate_something(self.btn_other, "Other"))

        self.to_find = QLineEdit(self)
        self.to_find.setMinimumHeight(35)
        self.to_find.setMaximumHeight(45)
        self.to_find.setStyleSheet("background-color: orange;")
        self.to_find.setFont(font)
        self.to_find.setPlaceholderText("Запрос...")

        self.btn_find = QPushButton("Найти", self)
        self.btn_find.setStyleSheet("background-color: orange;")
        self.btn_find.setMinimumHeight(35)
        self.btn_find.setMaximumHeight(45)
        self.btn_find.setFont(font)
        self.btn_find.clicked.connect(self.find_it)

        self.create_proj = QPushButton("Создать проект", self)
        self.create_proj.setMinimumHeight(35)
        self.create_proj.setMaximumHeight(45)
        self.create_proj.setStyleSheet("background-color: orange;")
        self.create_proj.setFont(font)
        self.create_proj.clicked.connect(self.create_project)

        self.delete_proj = QPushButton("Удалить проект", self)
        self.delete_proj.setMinimumHeight(35)
        self.delete_proj.setMaximumHeight(45)
        self.delete_proj.setStyleSheet("background-color: orange;")
        self.delete_proj.setFont(font)
        self.delete_proj.clicked.connect(self.delete_project)

        layout.addWidget(lbl, 0, 0, 25, 25)
        layout.addWidget(lbl3, 1, 0, 3, 5)
        layout.addWidget(self.btn_all, 3, 0, 3, 5)
        layout.addWidget(self.btn_it, 6, 0, 3, 5)
        layout.addWidget(self.btn_eng, 9, 0, 3, 5)
        layout.addWidget(self.btn_social, 12, 0, 3, 5)
        layout.addWidget(self.btn_other, 15, 0, 3, 5)
        layout.addWidget(self.to_find, 1, 5, 1, 5)
        layout.addWidget(self.btn_find, 1, 10, 1, 5)
        layout.addWidget(scroll_area, 2, 5, 21, 18)
        layout.addWidget(self.back_btn, 23, 0, 2, 5)
        layout.addWidget(self.create_proj, 23, 5, 2, 5)
        layout.addWidget(self.delete_proj, 23, 10, 2, 5)
        self.setLayout(layout)

    def create_project(self):
        if not authorized:
            QMessageBox.critical(self, "Авторизация", "Вы не авторизованы. Войдите в аккаунт в настройках.")
            return
        stack_root.setCurrentIndex(6)
        stack_root.setMinimumSize(create_proj_w.width, create_proj_w.height)
        stack_root.setMaximumSize(int(create_proj_w.width * 1.1), int(create_proj_w.height * 1.1))
        desktop_geometry = QGuiApplication.primaryScreen().availableGeometry()
        x = (desktop_geometry.width() - create_proj_w.width) // 2
        y = (desktop_geometry.height() - create_proj_w.height) // 2
        stack_root.setGeometry(x, y, create_proj_w.width, create_proj_w.height)

    def delete_project(self):
        global del_proj_w, result
        if not authorized:
            QMessageBox.critical(self, "Авторизация", "Вы не авторизованы. Войдите в аккаунт в настройках.")
            return

        try:
            response = requests.get(flask_server_url + "get_user_projs", data={"username": authorized[0]})
            response.raise_for_status()  # Проверка успешности запроса
        except requests.exceptions.RequestException as e:
            QMessageBox.critical(self, "Запрос", "Ошибка при подключении к серверу.")
            return

        if response.status_code == 200:
            json_data = response.json()
            if json_data:
                if json_data["status"] != "OK":
                    QMessageBox.critical(self, "Запрос", "Неизвестная ошибка.")
                    return
        else:
            QMessageBox.critical(self, "Запрос", f"Ошибка. Код состояния: {response.status_code}")
            return
        result = json_data["result"]
        stack_root.removeWidget(del_proj_w)
        del del_proj_w
        del_proj_w = SamDelProj(result)
        stack_root.insertWidget(7, del_proj_w)
        stack_root.setCurrentIndex(7)
        stack_root.setMinimumSize(del_proj_w.width, del_proj_w.height)
        stack_root.setMaximumSize(int(del_proj_w.width * 1.1), int(del_proj_w.height * 1.1))
        desktop_geometry = QGuiApplication.primaryScreen().availableGeometry()
        x = (desktop_geometry.width() - del_proj_w.width) // 2
        y = (desktop_geometry.height() - del_proj_w.height) // 2
        stack_root.setGeometry(x, y, del_proj_w.width, del_proj_w.height)

    def fly_bird(self):
        global authorized
        name = self.btn_names[self.sender()]
        if not authorized:
            QMessageBox.critical(self, "Авторизация", "Вы не авторизованы. Войдите в аккаунт в настройках.")
            return

        try:
            response = requests.post(flask_server_url + "subs_proj", json={"name": name, "username": authorized[0]})
            response.raise_for_status()  # Проверка успешности запроса
        except requests.exceptions.RequestException as e:
            QMessageBox.critical(self, "Запрос", "Ошибка при подключении к серверу.")
            return

        if response.status_code == 200:
            json_data = response.json()
            if json_data:
                if json_data["status"] == "OK":
                    QMessageBox.information(self, "Запрос", "Запрос на вступление был отправлен.")
                elif json_data["status"] == "error":
                    QMessageBox.warning(self, "Запрос", "Произошла ошибка.")
                else:
                    QMessageBox.critical(self, "Запрос", "Неизвестная ошибка.")
        else:
            QMessageBox.critical(self, "Запрос", f"Ошибка. Код состояния: {response.status_code}")

    def find_it(self):
        self.btn_all.setStyleSheet("background-color: orange;")
        self.btn_eng.setStyleSheet("background-color: #d0d0d0;")
        self.btn_it.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_social.setStyleSheet("background-color: #d0d0d0;")

        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        bird_pict = QIcon(QPixmap("Bird.png"))
        lst_wdgts = []
        self.btn_names = dict()
        for i in range(len(projects)):
            if self.to_find.text().lower() in projects[i][1].lower():
                name = QLineEdit(self)
                name.setText(projects[i][1])
                name.setFont(font)
                name.setMinimumHeight(35)
                name.setMaximumSize(800, 45)
                name.setReadOnly(True)
                name.setStyleSheet("background-color: #ffd550")

                bird = QPushButton(self)
                bird.setIcon(bird_pict)
                bird.setFixedSize(35, 35)
                bird.setIconSize(QSize(30, 30))
                bird.clicked.connect(self.fly_bird)
                self.btn_names[bird] = name.text()

                wdgt = QWidget(self)
                wdgt.setMaximumHeight(60)
                wdgt.setMinimumHeight(45)
                lyt = QHBoxLayout(wdgt)
                lyt.addWidget(name)
                lyt.addWidget(bird)
                lst_wdgts.append(wdgt)

                if projects[i][3]:
                    lbl2 = QLabel(self)
                    pixmap = QPixmap("Proj_pictures/" + projects[i][3])

                    lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                    lbl2.setPixmap(pixmap)
                    lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                    lst_wdgts.append(lbl2)

                text = QTextEdit(self)
                text.setPlainText(projects[i][2])
                text.setFont(font)
                text.setMinimumHeight(int(len(projects[i][2]) * 1.1))
                text.setReadOnly(True)
                lst_wdgts.append(text)

            for i in range(len(lst_wdgts)):
                self.layout2.addWidget(lst_wdgts[i])

    def activate_all(self):
        self.btn_all.setStyleSheet("background-color: orange;")
        self.btn_eng.setStyleSheet("background-color: #d0d0d0;")
        self.btn_it.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_social.setStyleSheet("background-color: #d0d0d0;")
        self.to_find.setText("")

        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        bird_pict = QIcon(QPixmap("Bird.png"))
        lst_wdgts = []
        self.btn_names = dict()
        for i in range(len(projects)):
            name = QLineEdit(self)
            name.setText(projects[i][1])
            name.setFont(font)
            name.setMinimumHeight(35)
            name.setMaximumSize(800, 45)
            name.setReadOnly(True)
            name.setStyleSheet("background-color: #ffd550")

            bird = QPushButton(self)
            bird.setIcon(bird_pict)
            bird.setFixedSize(35, 35)
            bird.setIconSize(QSize(30, 30))
            bird.clicked.connect(self.fly_bird)
            self.btn_names[bird] = name.text()

            wdgt = QWidget(self)
            wdgt.setMaximumHeight(60)
            wdgt.setMinimumHeight(45)
            lyt = QHBoxLayout(wdgt)
            lyt.addWidget(name)
            lyt.addWidget(bird)
            lst_wdgts.append(wdgt)

            if projects[i][3]:
                lbl2 = QLabel(self)
                pixmap = QPixmap("Proj_pictures/" + projects[i][3])

                lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                lbl2.setPixmap(pixmap)
                lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                lst_wdgts.append(lbl2)

            text = QTextEdit(self)
            text.setPlainText(projects[i][2])
            text.setFont(font)
            text.setMinimumHeight(int(len(projects[i][2]) * 1.1))
            text.setReadOnly(True)
            lst_wdgts.append(text)

        for i in range(len(lst_wdgts)):
            self.layout2.addWidget(lst_wdgts[i])

    def activate_something(self, wdgt, category):
        self.btn_it.setStyleSheet("background-color: #d0d0d0;")
        self.btn_eng.setStyleSheet("background-color: #d0d0d0;")
        self.btn_all.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_social.setStyleSheet("background-color: #d0d0d0;")
        wdgt.setStyleSheet("background-color: orange;")
        self.to_find.setText("")

        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        bird_pict = QIcon(QPixmap("Bird.png"))
        lst_wdgts = []
        self.btn_names = dict()
        for i in range(len(projects)):
            if projects[i][4] == category:
                name = QLineEdit(self)
                name.setText(projects[i][1])
                name.setFont(font)
                name.setMinimumHeight(35)
                name.setMaximumSize(800, 45)
                name.setReadOnly(True)
                name.setStyleSheet("background-color: #ffd550")

                bird = QPushButton(self)
                bird.setIcon(bird_pict)
                bird.setFixedSize(35, 35)
                bird.setIconSize(QSize(30, 30))
                bird.clicked.connect(self.fly_bird)
                self.btn_names[bird] = name.text()

                wdgt = QWidget(self)
                wdgt.setMaximumHeight(60)
                wdgt.setMinimumHeight(45)
                lyt = QHBoxLayout(wdgt)
                lyt.addWidget(name)
                lyt.addWidget(bird)
                lst_wdgts.append(wdgt)

                if projects[i][3]:
                    lbl2 = QLabel(self)
                    pixmap = QPixmap("Proj_pictures/" + projects[i][3])

                    lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                    lbl2.setPixmap(pixmap)
                    lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                    lst_wdgts.append(lbl2)

                text = QTextEdit(self)
                text.setPlainText(projects[i][2])
                text.setFont(font)
                text.setMinimumHeight(int(len(projects[i][2]) * 1.1))
                text.setReadOnly(True)
                lst_wdgts.append(text)

            for i in range(len(lst_wdgts)):
                self.layout2.addWidget(lst_wdgts[i])


class SamProfilesOrgs(QWidget):
    def __init__(self):
        super().__init__()
        self.width = 1200
        self.height = 700

        self.setMinimumSize(self.width, self.height)

        lbl = QLabel(self)
        pixmap = QPixmap("Cloth.jpg")
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        lbl.setScaledContents(True)  # Растягиваем содержимое QLabel

        lbl.setPixmap(pixmap)

        self.back_btn = QPushButton("Назад", self)
        self.back_btn.setMinimumHeight(35)
        self.back_btn.setMaximumHeight(45)
        self.back_btn.setStyleSheet("background-color: orange;")
        self.back_btn.setFont(font)
        self.back_btn.clicked.connect(lambda e: open_another_widget(0, (root.width, root.height),
                                                                    (int(root.width * 1.1), int(root.height * 1.1))))

        layout = QGridLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        # Создаем QScrollArea
        scroll_area = QScrollArea(self)
        scroll_area.setStyleSheet("background-color: #d0d0d0")
        scroll_area.setWidgetResizable(True)  # Разрешаем изменение размера виджета внутри QScrollArea

        # Создаем виджет, который будет помещен в QScrollArea
        scroll_content = QWidget(scroll_area)
        self.layout2 = QVBoxLayout(scroll_content)

        lst_wdgts = []
        for i in range(len(orgs)):
            name = QLineEdit(self)
            name.setText(orgs[i][1])
            name.setFont(font)
            name.setMinimumHeight(35)
            name.setMaximumHeight(45)
            name.setReadOnly(True)
            name.setStyleSheet("background-color: #ffd550")
            lst_wdgts.append(name)

            if orgs[i][6]:
                lbl2 = QLabel(self)
                pixmap = QPixmap("Prof_pictures/" + orgs[i][6])

                lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                lbl2.setPixmap(pixmap)
                lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                lst_wdgts.append(lbl2)

            text = QTextEdit(self)
            text.setPlainText(orgs[i][2])
            text.setFont(font)
            text.setMinimumHeight(int(len(orgs[i][2]) * 1.3))
            text.setReadOnly(True)
            lst_wdgts.append(text)

            chance = QTextEdit(self)
            chance.setPlainText("Активные проекты:\n\n" + orgs[i][5])
            chance.setFont(font)
            chance.setMinimumHeight(int(len(orgs[i][5]) * 1.3))
            chance.setReadOnly(True)
            lst_wdgts.append(chance)

        for i in range(len(lst_wdgts)):
            self.layout2.addWidget(lst_wdgts[i])

        # Устанавливаем виджет в QScrollArea
        scroll_area.setWidget(scroll_content)

        lbl3 = QLabel("Категории", self)
        lbl3.setMinimumHeight(35)
        lbl3.setMaximumHeight(45)
        lbl3.setStyleSheet("color: red;")
        lbl3.setFont(QFont("Segoe Print", 25))
        lbl3.setAlignment(Qt.AlignmentFlag.AlignCenter)

        self.btn_all = QPushButton("Все", self)
        self.btn_all.setStyleSheet("background-color: orange;")
        self.btn_all.setMinimumHeight(40)
        self.btn_all.setFont(font)
        self.btn_all.clicked.connect(self.activate_all)

        self.btn_children = QPushButton("Дети", self)
        self.btn_children.setStyleSheet("background-color: #d0d0d0;")
        self.btn_children.setMinimumHeight(40)
        self.btn_children.setFont(font)
        self.btn_children.clicked.connect(lambda e: self.activate_something(self.btn_children, "Children"))

        self.btn_help = QPushButton("Нуждающиеся", self)
        self.btn_help.setStyleSheet("background-color: #d0d0d0;")
        self.btn_help.setMinimumHeight(40)
        self.btn_help.setFont(font)
        self.btn_help.clicked.connect(lambda e: self.activate_something(self.btn_help, "Help"))

        self.btn_animals = QPushButton("Животные", self)
        self.btn_animals.setStyleSheet("background-color: #d0d0d0;")
        self.btn_animals.setMinimumHeight(40)
        self.btn_animals.setFont(font)
        self.btn_animals.clicked.connect(lambda e: self.activate_something(self.btn_animals, "Animals"))

        self.btn_other = QPushButton("Другое", self)
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setMinimumHeight(40)
        self.btn_other.setFont(font)
        self.btn_other.clicked.connect(lambda e: self.activate_something(self.btn_other, "Other"))

        self.to_find = QLineEdit(self)
        self.to_find.setMinimumHeight(35)
        self.to_find.setMaximumHeight(45)
        self.to_find.setStyleSheet("background-color: orange;")
        self.to_find.setFont(font)
        self.to_find.setPlaceholderText("Запрос...")

        self.btn_find = QPushButton("Найти", self)
        self.btn_find.setStyleSheet("background-color: orange;")
        self.btn_find.setMinimumHeight(35)
        self.btn_find.setMaximumHeight(45)
        self.btn_find.setFont(font)
        self.btn_find.clicked.connect(self.find_it)

        layout.addWidget(lbl, 0, 0, 25, 25)
        layout.addWidget(lbl3, 1, 0, 3, 5)
        layout.addWidget(self.btn_all, 3, 0, 3, 5)
        layout.addWidget(self.btn_children, 6, 0, 3, 5)
        layout.addWidget(self.btn_help, 9, 0, 3, 5)
        layout.addWidget(self.btn_animals, 12, 0, 3, 5)
        layout.addWidget(self.btn_other, 15, 0, 3, 5)
        layout.addWidget(self.to_find, 1, 5, 1, 5)
        layout.addWidget(self.btn_find, 1, 10, 1, 5)
        layout.addWidget(scroll_area, 2, 5, 21, 18)
        layout.addWidget(self.back_btn, 23, 0, 2, 5)
        self.setLayout(layout)

    def find_it(self):
        self.btn_all.setStyleSheet("background-color: orange;")
        self.btn_help.setStyleSheet("background-color: #d0d0d0;")
        self.btn_children.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_animals.setStyleSheet("background-color: #d0d0d0;")

        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        lst_wdgts = []
        for i in range(len(orgs)):
            if self.to_find.text().lower() in orgs[i][1].lower():
                name = QLineEdit(self)
                name.setText(orgs[i][1])
                name.setFont(font)
                name.setMinimumHeight(35)
                name.setMaximumHeight(45)
                name.setReadOnly(True)
                name.setStyleSheet("background-color: #ffd550")
                lst_wdgts.append(name)

                if orgs[i][6]:
                    lbl2 = QLabel(self)
                    pixmap = QPixmap("Prof_pictures/" + orgs[i][6])

                    lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                    lbl2.setPixmap(pixmap)
                    lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                    lst_wdgts.append(lbl2)

                text = QTextEdit(self)
                text.setPlainText(orgs[i][2])
                text.setFont(font)
                text.setMinimumHeight(int(len(orgs[i][2]) * 1.3))
                text.setReadOnly(True)
                lst_wdgts.append(text)

                chance = QTextEdit(self)
                chance.setPlainText("Активные проекты:\n\n" + orgs[i][5])
                chance.setFont(font)
                chance.setMinimumHeight(int(len(orgs[i][5]) * 1.3))
                chance.setReadOnly(True)
                lst_wdgts.append(chance)

        for i in range(len(lst_wdgts)):
            self.layout2.addWidget(lst_wdgts[i])

    def activate_all(self):
        self.btn_all.setStyleSheet("background-color: orange;")
        self.btn_help.setStyleSheet("background-color: #d0d0d0;")
        self.btn_children.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_animals.setStyleSheet("background-color: #d0d0d0;")
        self.to_find.setText("")

        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        lst_wdgts = []
        for i in range(len(orgs)):
            name = QLineEdit(self)
            name.setText(orgs[i][1])
            name.setFont(font)
            name.setMinimumHeight(35)
            name.setMaximumHeight(45)
            name.setReadOnly(True)
            name.setStyleSheet("background-color: #ffd550")
            lst_wdgts.append(name)

            if orgs[i][6]:
                lbl2 = QLabel(self)
                pixmap = QPixmap("Prof_pictures/" + orgs[i][6])

                lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                lbl2.setPixmap(pixmap)
                lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                lst_wdgts.append(lbl2)

            text = QTextEdit(self)
            text.setPlainText(orgs[i][2])
            text.setFont(font)
            text.setMinimumHeight(int(len(orgs[i][2]) * 1.3))
            text.setReadOnly(True)
            lst_wdgts.append(text)

            chance = QTextEdit(self)
            chance.setPlainText("Активные проекты:\n\n" + orgs[i][5])
            chance.setFont(font)
            chance.setMinimumHeight(int(len(orgs[i][5]) * 1.3))
            chance.setReadOnly(True)
            lst_wdgts.append(chance)

        for i in range(len(lst_wdgts)):
            self.layout2.addWidget(lst_wdgts[i])

    def activate_something(self, wdgt, category):
        self.btn_children.setStyleSheet("background-color: #d0d0d0;")
        self.btn_help.setStyleSheet("background-color: #d0d0d0;")
        self.btn_all.setStyleSheet("background-color: #d0d0d0;")
        self.btn_other.setStyleSheet("background-color: #d0d0d0;")
        self.btn_animals.setStyleSheet("background-color: #d0d0d0;")
        wdgt.setStyleSheet("background-color: orange;")
        self.to_find.setText("")

        while self.layout2.count():
            item = self.layout2.takeAt(0)
            widget = item.widget()
            if widget:
                widget.deleteLater()  # Освобождаем память для виджета
            else:
                del item  # Освобождаем память для элемента макета

        lst_wdgts = []
        for i in range(len(orgs)):
            if orgs[i][4] == category:
                name = QLineEdit(self)
                name.setText(orgs[i][1])
                name.setFont(font)
                name.setMinimumHeight(35)
                name.setMaximumHeight(45)
                name.setReadOnly(True)
                name.setStyleSheet("background-color: #ffd550")
                lst_wdgts.append(name)

                if orgs[i][6]:
                    lbl2 = QLabel(self)
                    pixmap = QPixmap("Prof_pictures/" + orgs[i][6])

                    lbl2.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    lbl2.setScaledContents(True)  # Растягиваем содержимое QLabel
                    lbl2.setPixmap(pixmap)
                    lbl2.setMaximumSize(670, int(pixmap.height() * 670 / pixmap.width()))
                    lst_wdgts.append(lbl2)

                text = QTextEdit(self)
                text.setPlainText(orgs[i][2])
                text.setFont(font)
                text.setMinimumHeight(int(len(orgs[i][2]) * 1.3))
                text.setReadOnly(True)
                lst_wdgts.append(text)

                chance = QTextEdit(self)
                chance.setPlainText("Активные проекты:\n\n" + orgs[i][5])
                chance.setFont(font)
                chance.setMinimumHeight(int(len(orgs[i][5]) * 1.3))
                chance.setReadOnly(True)
                lst_wdgts.append(chance)

        for i in range(len(lst_wdgts)):
            self.layout2.addWidget(lst_wdgts[i])


def open_another_widget(index, min_sizes, max_sizes, *, state=None):
    stack_root.setCurrentIndex(index)
    stack_root.setMinimumSize(min_sizes[0], min_sizes[1])
    stack_root.setMaximumSize(max_sizes[0], max_sizes[1])
    desktop_geometry = QGuiApplication.primaryScreen().availableGeometry()
    x = (desktop_geometry.width() - min_sizes[0]) // 2
    y = (desktop_geometry.height() - min_sizes[1]) // 2
    stack_root.setGeometry(x, y, min_sizes[0], min_sizes[1])


if __name__ == '__main__':
    app = QApplication(sys.argv)

    root = SamRoot()
    scalendar = SamCalendar()
    diary_builder = SamDiaryBuilder()
    world_projects = SamWorldProjects()
    profiles_orgs = SamProfilesOrgs()
    settings = SamSettings()
    create_proj_w = SamNewProj()
    del_proj_w = SamDelProj([])

    stack_root = SamStack(root, scalendar, diary_builder, world_projects, profiles_orgs, settings, create_proj_w,
                          del_proj_w)
    stack_root.show()

    sys.exit(app.exec())
